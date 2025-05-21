# vector_store.py
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.docstore.document import Document
from docx import Document as DocxDocument
from dotenv import load_dotenv
import os, hashlib

# Загрузка .env вручную (обязательно для сервисов)
load_dotenv(dotenv_path="/opt/dzvfr/.env")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

REPORTS_DIR = "reports"
INDEX_PATH = "data/faiss_index"

# Простой хеш по содержимому файла для избежания повторов
def file_hash(filepath):
    with open(filepath, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()

# Парсинг .docx в список чанков с метаинфой
def parse_docx(filepath):
    doc = DocxDocument(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    filename = os.path.basename(filepath)
    docs = []
    for i, para in enumerate(paragraphs):
        docs.append(Document(
            page_content=para,
            metadata={"source": filename, "chunk": i}
        ))
    return docs

# Загрузка всех документов из папки
def load_all_docs():
    all_docs = []
    seen = set()
    for file in os.listdir(REPORTS_DIR):
        if file.endswith(".docx"):
            path = os.path.join(REPORTS_DIR, file)
            h = file_hash(path)
            if h not in seen:
                seen.add(h)
                all_docs.extend(parse_docx(path))
    return all_docs

# Инициализация или загрузка FAISS
def init_vector_store():
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_API_KEY)
    if os.path.exists(INDEX_PATH):
        return FAISS.load_local(INDEX_PATH, embeddings, allow_dangerous_deserialization=True)

    docs = load_all_docs()
    db = FAISS.from_documents(docs, embeddings)
    db.save_local(INDEX_PATH)
    return db
