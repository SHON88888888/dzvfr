# parser.py
from docx import Document as DocxDocument
from langchain.docstore.document import Document
import os

# Парсим .docx, сохраняем метаинфу: файл, номер чанка, возможно раздел (если найден)
def parse_docx(filepath):
    doc = DocxDocument(filepath)
    filename = os.path.basename(filepath)

    sections = []
    current_section = "Общий раздел"
    buffer = []
    results = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Простейшее определение раздела — заглавные, жирные, центрованные
        if p.style.name.startswith("Heading") or (p.alignment == 1 and text.isupper()):
            # Сохраняем предыдущий блок
            if buffer:
                combined = " ".join(buffer)
                results.append(Document(
                    page_content=combined,
                    metadata={"source": filename, "section": current_section}
                ))
                buffer = []
            current_section = text
        else:
            buffer.append(text)

    # Финальный блок
    if buffer:
        combined = " ".join(buffer)
        results.append(Document(
            page_content=combined,
            metadata={"source": filename, "section": current_section}
        ))

    return results
